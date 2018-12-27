from docx import Document
import os
import xlrd
import difflib
from xlutils.copy import copy

document = Document('参考答案.docx')
answer_dict = {}
normal_answer_list = []
count = 1

for paragraph in document.paragraphs:
    is_answer = False
    answer = ''
    para_list = paragraph.text.split()
    for item in para_list:
        if '.' in list(item):
            is_answer = True
    if len(para_list) and is_answer:
        answer_list = paragraph.text.split()
        new_list = []
        replace_flag = False
        for i in range(len(answer_list)):
            answer_list[i] = answer_list[i].replace('.', ' ').strip().split()
            if len(answer_list[i]) > 1:
                replace_flag = True
                for j in range(len(answer_list[i])):
                    temp = [answer_list[i][j]]
                    new_list.append(temp)
            else:
                new_list.append(answer_list[i])
        if replace_flag:
            answer_list = new_list
        for i in range(len(answer_list)):
            item = answer_list[i][0]
            if item == str(count):
                count += 1
                if answer:
                    answer_dict[count - 2] = answer.strip()
                    normal_answer_list.append(answer.strip())
                    answer = ''
            else:
                answer = answer + item + ' '
                if i == len(answer_list) - 1:
                    answer_dict[count - 1] = answer.strip()
                    normal_answer_list.append(answer.strip())

answer_dict[33] = answer_dict[33].replace(' ', '.')
for i in range(35, 40):
    answer_dict[i + 1] = answer_dict[i + 1].replace(', ', '')

files = os.listdir('答题卡')
for file in files:
    file_num = files.index(file)
    grade = [0, 0, 0, 0]
    stu_answer = []
    if not os.path.isdir(file):
        stu_document = Document("答题卡/" + file)
        tables = stu_document.tables
        for num in range(1, 3):
            table = tables[num]
            for i in range(len(table.rows)):
                if i % 2 == 1:
                    for j in range(len(table.columns)):
                        stu_answer.append(table.cell(i, j).text)
        for num in range(3, 5):
            table = tables[num]
            for i in range(len(table.rows)):
                for j in range(len(table.columns)):
                    if j == 1:
                        stu_answer.append(table.cell(i, j).text)
        table = tables[5]
        stu_answer.append(table.cell(0, 0).text)

        table = tables[0]
        for i in range(20):
            if i == 17:
                if (stu_answer[i].lower() == 'improvements' or
                        stu_answer[i].lower() == 'improvement' or
                        stu_answer[i].lower() == 'improving'):
                    grade[0] += 1
            elif i == 18:
                if (stu_answer[i].lower() == 'will meet' or
                        stu_answer[i].lower() == 'will be meeting'):
                    grade[0] += 1
            elif stu_answer[i].lower() == answer_dict[i + 1].lower():
                grade[0] += 1
        table.cell(1, 1).text = str(grade[0])
        for i in range(20, 45):
            if stu_answer[i].lower() == answer_dict[i + 1].lower():
                grade[1] += 2
        table.cell(1, 2).text = str(grade[1])
        for i in range(45, 49):
            if stu_answer[i].lower() == answer_dict[i + 1].lower():
                grade[2] += 2
        sentence = stu_answer[49]
        if sentence:
            rate = difflib.SequenceMatcher(None, sentence, answer_dict[
                50]).ratio()
            if rate >= 0.6:
                grade[2] += 7
            elif 0.3 < rate < 0.6:
                grade[2] += 6
        else:
            pass
        table.cell(1, 3).text = str(grade[2])

        article = stu_answer[50]
        if article:
            word_count = len(article.split())
            if word_count in range(20, 30):
                grade[3] = 6
            elif word_count in range(30, 40):
                grade[3] = 8
            elif word_count in range(40, 50):
                grade[3] = 10
            elif word_count in range(50, 60):
                grade[3] = 12
            elif word_count >= 60:
                grade[3] = 13
        else:
            pass
        table.cell(1, 4).text = str(grade[3])
        table.cell(1, 5).text = str(grade[0] + grade[1] + grade[2] + grade[3])

        excel = xlrd.open_workbook('成绩单.xls')
        wb = copy(excel)
        ws = wb.get_sheet(0)
        ws.write(file_num + 1, 2, grade[0] + grade[1] + grade[2] + grade[3])
        wb.save('成绩单.xls')

        stu_document.save(r"答题卡/" + file)
